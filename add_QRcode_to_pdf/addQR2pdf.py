# !/usr/bin/env python
# -*-coding:utf-8 -*-

"""
# File       : addQR2pdf.py
# Time       ：6/9/2023 22:28
# Author     ：Allen Wong
# version    ：python 3.10
# Description：读取PDF文件，产生一个新的PDF；在第一个页面固定位置（）写入二维码
"""

import PyPDF2
import qrcode
from docx import Document
from docx.shared import Inches, Cm
from docx2pdf import convert  # 需要安装office才可使用这个包
from PyPDF2 import PdfFileMerger
import PySimpleGUI as sg  # 图形选择窗口
import os


# 产生一个图形窗口，供用户选择要添加的PDF文件
def create_pdf_select_gui() -> str:
    file_path = ''

    # define layout
    layout = [
        [sg.Text('Select the PDF file to add QR:')],
        [sg.InputText(disabled=True, do_not_clear=False), sg.FileBrowse()],
        [sg.Text('Choose Jacket:')],
        [sg.DropDown(['Jacket1', 'Jacket2', 'Jacket3'], key='-DROPDOWN-')],
        [sg.Submit()]
    ]

    window = sg.Window('CBIC QRcode Stamping Gadget', layout)
    # 获取事件

    while True:
        event, values = window.read()

        match event:
            case sg.WINDOW_CLOSED:
                break
            case "Submit":
                file_path = values[0]
                if not file_path.endswith(".pdf") and not file_path.endswith(".PDF"):
                    sg.popup('Please choose .pdf or .PDF extension file.', title='Error Message')
                else:
                    select_opt = values['-DROPDOWN-']  # select_opt 是字符串
                    if select_opt == '' or select_opt is None:
                        sg.popup('Please choose one dropdown option', title='Error Message')
                    else:
                        break
            case _:
                sg.popup("Unknown error.")

    window.close()

    return file_path,select_opt


# statistic ver1
def genQR(qrcontent):

    qrimg = qrcode.make(qrcontent)
    qrimg.save('./qrcode.jpg')
    img = './qrcode.jpg'
    return (img)


# 将二维码放入WORD
def replaceWord(img, TEMPDOC, TEMP_REPLACE_DOC):
    doc = Document(TEMPDOC)
    for paragrah in doc.paragraphs:
        if '<<imgX>>' in paragrah.text:
            paragrah.text = '     ' + paragrah.text.replace('<<imgX>>', ' ')
            run = paragrah.add_run('')
            run.add_picture(img, width=Cm(3), height=None)
    doc.save(TEMP_REPLACE_DOC)


# 使用docx2pdf包
def convert_to_pdf(TEMP_REPLACE_DOC, TEMP_REPLACE_PDF):
    convert('./temp_replace.docx', './temp_replace.pdf')
    # convert(TEMP_REPLACE_DOC, TEMP_REPLACE_PDF)
    return './temp_replace.pdf'

def choose_Jacket(select_opt) -> str:
    match select_opt:
        case 'Jacket1':
            return 'QRCODE - JACKET1'
        case 'Jacket2':
            return 'QRCODE - JACKET2'
        case 'Jacket3':
            return 'QRCODE - JACKET3'
        case _:
            print('JACKET CHOOSE ERROR')
            exit(1)

# 使用其他方式实现docx2pdf
# def convert_to_pdf(TEMP_REPLACE_DOC, TEMP_REPLACE_PDF):


# 合并temp_replace.pdf和第一页

def main(addQRonpage=None):

    print('START.')

    # 常量
    TEMPDOC = './temp.docx'  # 模版word文件，用来里面有个标识<<imgX>>用来表示替换的位置
    TEMP_REPLACE_DOC = './temp_replace.docx'  # 生成带二维码的word文件
    TEMP_REPLACE_PDF = './temp_replace.pdf'  # 中间临时的PDF

    # UI界面确定选择PDF文件
    input_pdf,select_opt = create_pdf_select_gui()

    # 确定产生哪个Jacket

    qrcontent = choose_Jacket(select_opt)

    # 产生二维码

    img = genQR(qrcontent)

    # 替换二维码，生成临时的文件 './temp_replace.docx'

    replaceWord(img, TEMPDOC, TEMP_REPLACE_DOC)

    # 将替换后的文件转换成PDF

    convert_to_pdf(TEMP_REPLACE_DOC, TEMP_REPLACE_PDF)

    # 读取INPUT_PDF并且将二维码合并在第一个页面上

    with open(input_pdf, 'rb') as file:
        # read file
        pdf_reader = PyPDF2.PdfReader(file)
        pdf_writer = PyPDF2.PdfWriter()  # 创建一个写入器

        # for page_num in range(pdf_reader.numPages): # Mac version
        for page_num in range(len(pdf_reader.pages)): # Win version
            # page = pdf_reader.getPage(page_num) # Mac version
            page = pdf_reader.pages[page_num] # Win version
            # 对第一页进行修改
            if page_num == 0:
                SignPDF = './temp_replace.pdf'
                pdf_readerX = PyPDF2.PdfReader(SignPDF)
                # page.mergePage(pdf_readerX.getPage(0))  # 此函数重要，mergePage（另一个PDF页面） Mac version
                page.merge_page(pdf_readerX.pages[0])  # 此函数重要，mergePage（另一个PDF页面） Win Version
            # 其他页
            pdf_writer.add_page(page) #Mac Version

    # write file
    addQRfile = './' + input_pdf.rsplit('/')[-1].split('.')[0] + '_QR.pdf'

    with open(addQRfile, 'wb') as output_file:
        pdf_writer.write(output_file)

    print('END.')


if __name__ == '__main__':
    main()