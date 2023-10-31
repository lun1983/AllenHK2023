# !/usr/bin/env python
# -*-coding:utf-8 -*-

"""
# File       : create_doc.py
# Time       ：29/10/2023 21:37
# Author     ：Allen Wong
# version    ：python 3.11
# Conda Env  : CBIC v23.9.0
# Description：
"""

# %%
import os
import sys
from docx import Document
from docx.shared import Inches
from docx2pdf import convert


def create_doc_with_bar(bar_path: str, save_path: str, pdf_path: str, bar_valid_date: str):
    BOCOM_LOGO = os.path.join(os.path.dirname(sys.executable), "bocom_logo.jpg") if getattr(sys, 'frozen',
                                                                                            False) else "bocom_logo.jpg"
    TEMPLATE_PATH = os.path.join(os.path.dirname(sys.executable), "template_new.docx") if getattr(sys, 'frozen',
                                                                                                  False) else "template_new.docx"
    doc = Document(TEMPLATE_PATH)

    bar_valid_date_loc = doc.sections[0].header.paragraphs[1]
    # bar_valid_date_loc.text = bar_valid_date_loc.text.replace("XXXXXXXXXXXXXXX", bar_valid_date)

    replace_word = 'DDMMYYYY'

    # 表格内容的替换
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if replace_word in paragraph.text:
                        # paragraph.text = paragraph.text.replace(replace_word, bar_valid_date)
                        runs = paragraph.runs
                        for j in range(len(runs)):
                            if replace_word in runs[j].text:
                                runs[j].text = runs[j].text.replace(replace_word, bar_valid_date)

    # 非表格内容的替换
    for paragrah in doc.paragraphs:
        if replace_word in paragrah.text:

            # 不保留原格式替换
            # paragrah.text = paragrah.text.replace(replace_word, bar_valid_date)

            # 保持原格式替换
            runs = paragrah.runs
            for i in range(len(runs)):
                if replace_word in runs[i].text:
                    runs[i].text = runs[i].text.replace(replace_word, bar_valid_date)

    # bocom_logo_table = doc.sections[0].header.tables[0]
    # p = bocom_logo_table.rows[0].cells[0].add_paragraph()
    # r = p.add_run()
    # r.add_picture(BOCOM_LOGO,width=Inches(3), height=Inches(0.85))

    barcode_table = doc.tables[0].rows[2].cells[0].tables[0]
    p = barcode_table.rows[0].cells[0].add_paragraph()
    r = p.add_run()
    p.alignment = 1
    r.add_picture(bar_path, width=Inches(4.05), height=Inches(3.15))
    doc.save(save_path)
    convert(save_path, pdf_path)
