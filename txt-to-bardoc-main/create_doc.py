#%%
import os
import sys
from docx import Document
from docx.shared import Inches
from docx2pdf import convert



def create_doc_with_bar(bar_path: str, save_path: str, pdf_path: str, bar_valid_date: str):
    BOCOM_LOGO = os.path.join(os.path.dirname(sys.executable), "bocom_logo.jpg") if getattr(sys, 'frozen', False) else "bocom_logo.jpg"
    TEMPLATE_PATH = os.path.join(os.path.dirname(sys.executable), "template_new.docx") if getattr(sys, 'frozen', False) else "template_new.docx"
    doc = Document(TEMPLATE_PATH)
    
    bar_valid_date_loc = doc.sections[0].header.paragraphs[1]
    # bar_valid_date_loc.text = bar_valid_date_loc.text.replace("XXXXXXXXXXXXXXX", bar_valid_date)

    for paragrah in doc.paragraphs:
        if 'DD/MM/YYYY' in paragrah.text:
            paragrah.text = '     ' + paragrah.text.replace('DD/MM/YYYY', 'bar_valid_date')

    # bocom_logo_table = doc.sections[0].header.tables[0]
    # p = bocom_logo_table.rows[0].cells[0].add_paragraph()
    # r = p.add_run()
    # r.add_picture(BOCOM_LOGO,width=Inches(3), height=Inches(0.85))
    
    barcode_table = doc.tables[0].rows[2].cells[0].tables[0]
    p = barcode_table.rows[0].cells[0].add_paragraph()
    r = p.add_run()
    p.alignment = 1
    r.add_picture(bar_path,width=Inches(2.5), height=Inches(2.5))
    doc.save(save_path)
    print('save_path:'+str(save_path))
    print('pdf_path:'+str(pdf_path))
    convert(save_path, pdf_path)
    print('hello2')

