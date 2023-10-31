"""
# File       : addQR2pdf.py
# Time       ：19/9/2023 22:28
# Author     ：Allen Wong
# version    ：python 3.10
# Description：检查word中是否有包含关键字，如果有，则输出到 check_result.txt中
"""

from docx import Document
import os

CHECK_WORDS = []
CHECK_WORDS_FILE = './checkwordings.txt'
CHECK_FILE = []
CHECK_FILE_REULST = './check_result.txt'
CHECK_RESULT = []
CNT = 0


def readCheckwordingsfile(CHECK_WORDS_FILE):
    with open(CHECK_WORDS_FILE, 'r', encoding='utf-8') as file:
        for line in file.readlines():
            CHECK_WORDS.append(line.replace('\n', ''))

    return


def readFile(CHECK_FILE, READ_FILE):
    # 每次读文件之前，先清空列表
    # READ_FILE = []
    doc = Document(CHECK_FILE)
    # 读取纯文本
    for paragraph in doc.paragraphs:
        READ_FILE.append(paragraph.text)
    # 读取表格
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                # row_data.append(cell_text)
                READ_FILE.append(cell_text)

    return READ_FILE


def checkWord(checkword, CNT, READ_FILE):
    for item in READ_FILE:
        if checkword in item:
            CNT = CNT + 1
            print(str(CNT) + ' Checkword Found: ' + checkword)
            CHECK_RESULT.append(str(CNT) + ' Checkword Found: ' + checkword)

    return CNT


def wirte_output(CHECK_RESULT):
    with open(CHECK_FILE_REULST, 'w', encoding='utf-8') as file:
        for resultx in CHECK_RESULT:
            file.write(resultx + '\n')

    return


def readCheckfileName():
    for filename in os.listdir('./'):
        if filename.endswith('.doc') or filename.endswith('.docx'):
            CHECK_FILE.append('./' + filename)

    return


def main():
    # 检查字典
    if not os.path.exists(CHECK_WORDS_FILE):
        print('File:\'checkwordings.txt\' not exists in current path')
        exit(1)
    # 将checkwordings文件读到列表
    readCheckwordingsfile(CHECK_WORDS_FILE)

    # 当前文件夹下的所有doc和docx文件名读到 CHECK_FILE[]
    readCheckfileName()

    READ_FILE = []
    for file in CHECK_FILE:
        checkEachFile(file, READ_FILE)
    # 将结果写到txt文件
    wirte_output(CHECK_RESULT)

    return


def checkEachFile(CHECK_FILE, READ_FILE):
    # 读取要检查的文件
    print('*****Now checking: ' + CHECK_FILE + '*****')
    CHECK_RESULT.append('*****Now checking: ' + CHECK_FILE + '*****')
    READ_FILE.clear()

    # read file
    readFile(CHECK_FILE, READ_FILE)

    # 用CHECK_WORDS 的单词去匹配列表中的每一行，如果匹配到，则写入到 CHECK_RESULT
    CNT = 0

    for checkword in CHECK_WORDS:
        CNT = checkWord(checkword, CNT, READ_FILE)

    return


if __name__ == '__main__':
    main()
